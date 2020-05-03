import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_index_table(document, num_prac):       #Function to create Index table
    document = Document() 

    para = document.add_paragraph()

    para.alignment= WD_ALIGN_PARAGRAPH.CENTER #The Word "Index"
    run = para.add_run("INDEX\n")
    font = run.font
    font.name = "Arial"
    font.bold = True
    font.size = Pt(20)

    table = document.add_table(rows=num_prac+1, cols=3)
    table.cell(0,0).text="S.No"
    table.cell(0,1).text="Program"
    table.cell(0,2).text="Date"

    for i in range(1,num_prac+1):
        table.cell(i,0).text=f"{i}"

    document.add_page_break() # To to move to the next page
    return document
    

def create_exp(document, num_prac):           #Function to create Experiment page layout
     
    for i in range(1,num_prac+1):

        para1 = document.add_paragraph() #Aim of the experiment

        para1.alignment= WD_ALIGN_PARAGRAPH.LEFT
        run1 = para1.add_run(f"{i}. Enter the aim of the experiment\n")
        font1 = run1.font
        font1.name = "Arial"
        font1.bold = True
        font1.size = Pt(20)

        para2 = document.add_paragraph()

        para2.alignment= WD_ALIGN_PARAGRAPH.CENTER #The Word "Code"
        run2 = para2.add_run("CODE\n")
        font2 = run2.font
        font2.name = "TimesNewRoman"
        font2.bold = True
        font2.size = Pt(15)

        para3 = document.add_paragraph()

        para3.alignment= WD_ALIGN_PARAGRAPH.LEFT
        run3 = para3.add_run("Select this piece of text and Paste your code here ,with option unformated(Libre Office) or merge formating(Microsoft Word) \n\n")
        font3 = run3.font
        font3.name = "Calibri"
        font3.size = Pt(12)

        para4 = document.add_paragraph()

        para4.alignment= WD_ALIGN_PARAGRAPH.CENTER #The Word "OUTPUT"
        run4 = para4.add_run("OUTPUT\n")
        font4 = run4.font
        font4.name = "TimesNewRoman"
        font4.bold = True
        font4.size = Pt(15)

        document.add_page_break()
    
    return document

    
def main():
    print("PRACTICAL FILE TEMPLATE CREATOR")
    num_prac= int(input("\n\nEnter the number of experiments: "))
    document = Document()
    document = create_index_table(document, num_prac)
    document = create_exp(document, num_prac)
   
    document.save(os.path.dirname(__file__) + "/practical_file.docx")
    

main()
